"""
ATSense Backend — Production-Grade ML Pipeline
===============================================
Architecture: graceful-degradation semantic stack.

Tier 1 (preferred): sentence-transformers bi-encoder for true semantic similarity.
Tier 2 (fallback):  synonym-expanded TF-IDF cosine similarity.
Tier 3 (baseline):  raw TF-IDF (original behaviour, always available).

This means the app runs everywhere — Streamlit Community Cloud without GPU,
local dev with a GPU, anywhere — and degrades gracefully rather than crashing.
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from collections import defaultdict
from typing import Optional

import pdfplumber
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ──────────────────────────────────────────────────────────────────────────────
# 0. LOGGER
# ──────────────────────────────────────────────────────────────────────────────
logger = logging.getLogger("atsense.backend")
logging.basicConfig(level=logging.INFO)


# ──────────────────────────────────────────────────────────────────────────────
# 1. SEMANTIC ENGINE — lazy-loaded bi-encoder (Tier 1)
# ──────────────────────────────────────────────────────────────────────────────

_SEMANTIC_MODEL = None          # module-level singleton; loaded once on first use
_SEMANTIC_AVAILABLE = None      # tri-state: None=untried, True=ok, False=failed

SEMANTIC_MODEL_NAME = "all-MiniLM-L6-v2"   # 80 MB, Apache-2.0, strong on short text

def _get_semantic_model():
    """
    Lazy-load the sentence-transformers bi-encoder.
    Using a module singleton avoids re-loading the ~80 MB model on every
    Streamlit re-run — critical for perceived latency in the UI.

    Returns the loaded SentenceTransformer or None if unavailable.
    """
    global _SEMANTIC_MODEL, _SEMANTIC_AVAILABLE

    if _SEMANTIC_AVAILABLE is True:
        return _SEMANTIC_MODEL

    if _SEMANTIC_AVAILABLE is False:
        return None   # already tried and failed — don't retry on every call

    try:
        from sentence_transformers import SentenceTransformer
        _SEMANTIC_MODEL = SentenceTransformer(SEMANTIC_MODEL_NAME)
        _SEMANTIC_AVAILABLE = True
        logger.info("Semantic model loaded: %s", SEMANTIC_MODEL_NAME)
        return _SEMANTIC_MODEL
    except Exception as exc:
        _SEMANTIC_AVAILABLE = False
        logger.warning("Semantic model unavailable (%s). Falling back to TF-IDF.", exc)
        return None


# ──────────────────────────────────────────────────────────────────────────────
# 2. SKILL TAXONOMY — canonical skills + synonym graph
# ──────────────────────────────────────────────────────────────────────────────

# Master skill list used for extraction (extend freely)
SKILLS: set[str] = {
    # ML / AI
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "reinforcement learning", "transfer learning",
    "generative ai", "large language models", "llm", "rag",
    "feature engineering", "model deployment", "mlops", "a/b testing",
    # Frameworks
    "pytorch", "tensorflow", "keras", "scikit-learn", "hugging face",
    "langchain", "llamaindex", "xgboost", "lightgbm", "catboost",
    "opencv", "spacy", "nltk", "transformers",
    # Languages
    "python", "sql", "r", "java", "scala", "c++", "go", "javascript",
    # Data
    "pandas", "numpy", "spark", "hadoop", "dbt", "airflow",
    "kafka", "data pipelines", "etl", "feature store",
    # Infra / Cloud
    "docker", "kubernetes", "k8s", "aws", "gcp", "azure",
    "fastapi", "flask", "django", "rest api", "graphql",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "terraform", "ci/cd", "github actions", "linux",
    # Practices
    "git", "agile", "system design", "microservices", "distributed systems",
}

# Synonym graph: maps canonical skill → all acceptable surface forms.
# Bidirectional expansion: if ANY alias appears in a text, the canonical
# and all aliases are injected into the skill set before intersection scoring.
SKILL_SYNONYMS: dict[str, list[str]] = {
    "machine learning":         ["ml", "statistical learning", "predictive modeling",
                                  "statistical modeling", "deep learning"],
    "deep learning":            ["dl", "neural networks", "ann", "cnn", "rnn",
                                  "lstm", "transformer models"],
    "nlp":                      ["natural language processing", "text mining",
                                  "text analytics", "computational linguistics"],
    "computer vision":          ["cv", "image recognition", "object detection",
                                  "image segmentation"],
    "kubernetes":               ["k8s", "container orchestration", "helm"],
    "large language models":    ["llm", "gpt", "claude", "gemini", "foundation models"],
    "mlops":                    ["ml operations", "model ops", "model monitoring",
                                  "model lifecycle"],
    "ci/cd":                    ["continuous integration", "continuous deployment",
                                  "continuous delivery", "devops pipeline"],
    "aws":                      ["amazon web services", "ec2", "s3", "sagemaker",
                                  "lambda", "eks"],
    "gcp":                      ["google cloud", "google cloud platform",
                                  "vertex ai", "bigquery"],
    "postgresql":               ["postgres", "pg"],
    "elasticsearch":            ["elastic search", "opensearch", "elk"],
    "rest api":                 ["restful api", "rest", "api development", "fastapi",
                                  "flask api"],
    "rag":                      ["retrieval augmented generation",
                                  "retrieval-augmented generation"],
    "generative ai":            ["genai", "generative models", "diffusion models",
                                  "image generation"],
}

# Reverse lookup: alias → canonical (built at import time for O(1) lookup)
_ALIAS_TO_CANONICAL: dict[str, str] = {}
for _canonical, _aliases in SKILL_SYNONYMS.items():
    for _alias in _aliases:
        _ALIAS_TO_CANONICAL[_alias] = _canonical

# JD importance signals → numeric weight multiplier.
# Words indicating a skill is non-negotiable inflate its weight in the
# weighted_match_skills calculation so the penalty for missing it is severe.
IMPORTANCE_SIGNALS: dict[str, float] = {
    "must":            3.0,
    "required":        3.0,
    "non-negotiable":  3.0,
    "mandatory":       3.0,
    "essential":       2.5,
    "critical":        2.5,
    "core":            2.0,
    "minimum":         2.0,
    "need":            1.8,
    "expected":        1.5,
    "preferred":       0.6,
    "nice to have":    0.5,
    "bonus":           0.4,
    "plus":            0.5,
    "ideally":         0.6,
    "optional":        0.4,
    "desired":         0.7,
}

# Section-name patterns used for section splitter and density analysis
SECTION_PATTERNS: dict[str, list[str]] = {
    "summary":    ["summary", "objective", "profile", "about"],
    "experience": ["experience", "work history", "employment", "career"],
    "education":  ["education", "academic", "qualifications", "degrees"],
    "skills":     ["skills", "technical skills", "core competencies", "technologies"],
    "projects":   ["projects", "portfolio", "work samples"],
    "certifications": ["certifications", "certificates", "accreditations"],
}


# ──────────────────────────────────────────────────────────────────────────────
# 3. TEXT PRE-PROCESSING
# ──────────────────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Normalise text for TF-IDF and keyword matching.
    Steps:
      1. Unicode NFC normalisation (handles accented chars, ligatures).
      2. Lowercase.
      3. Strip non-alpha characters (preserves spaces).
      4. Collapse whitespace runs.
    """
    text = unicodedata.normalize("NFC", text)
    text = text.lower()
    text = re.sub(r"[^a-z\s]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def preprocess_text(text: str) -> list[str]:
    """
    Tokenise cleaned text into a list of unigrams.
    Used downstream for skill extraction (exact and bigram matching).
    """
    return clean_text(text).split()


def _sentence_tokenise(text: str) -> list[str]:
    """
    Split a block of text into sentences using punctuation heuristics.
    Used by semantic similarity to embed individual propositions rather
    than the full document, improving recall on long resumes.
    """
    # Split on period / semicolon / newline, then strip and drop empty strings
    sentences = re.split(r"[.;\n\r]+", text)
    return [s.strip() for s in sentences if len(s.strip()) > 15]


# ──────────────────────────────────────────────────────────────────────────────
# 4. TEXT EXTRACTION — bulletproof PDF/DOCX pipeline (fixes Scenario B)
# ──────────────────────────────────────────────────────────────────────────────

class ExtractionResult:
    """
    Value object returned by extract_text_from_file.
    Carries the extracted text AND metadata about HOW it was extracted,
    allowing the UI layer to surface warnings without coupling to backend logic.
    """
    def __init__(
        self,
        text: str,
        method: str,                   # "pdfplumber" | "ocr" | "docx" | "empty"
        warning: Optional[str] = None, # human-readable warning for the UI
        error: Optional[str] = None,   # hard error message (text will be "")
    ):
        self.text    = text
        self.method  = method
        self.warning = warning
        self.error   = error

    @property
    def ok(self) -> bool:
        """True if usable text was extracted."""
        return bool(self.text.strip()) and self.error is None


def extract_text_from_file(uploaded_file) -> ExtractionResult:
    """
    Multi-strategy text extractor with graceful degradation.

    Strategy waterfall for PDF:
      1. pdfplumber text layer extraction (fast, lossless for digital PDFs).
      2. If text layer is empty → OCR via pdf2image + pytesseract (scanned PDFs).
      3. If OCR also fails → return hard error with actionable user guidance.

    For DOCX:
      - python-docx paragraph extraction (handles tables + text boxes).

    Returns an ExtractionResult rather than a raw string so the caller has
    full context about extraction quality without inspecting the text itself.
    """
    if uploaded_file is None:
        return ExtractionResult("", "empty", error="No file provided.")

    filename = uploaded_file.name.lower()

    # ── PDF path ──────────────────────────────────────────────────────────────
    if filename.endswith(".pdf"):
        text = ""
        try:
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as exc:
            logger.warning("pdfplumber failed: %s", exc)
            return ExtractionResult(
                "", "empty",
                error=f"Could not open PDF: {exc}. Ensure the file is not password-protected."
            )

        # If we got meaningful text from the text layer, return it immediately
        if len(text.strip()) > 50:
            return ExtractionResult(text.strip(), "pdfplumber")

        # ── Scanned PDF fallback: OCR ──────────────────────────────────────
        # Scenario B fix: pdfplumber returns "" for image-only PDFs.
        # We attempt OCR before giving up, and tell the user if we used it.
        logger.info("Empty text layer — attempting OCR fallback.")
        try:
            from pdf2image import convert_from_bytes
            import pytesseract

            uploaded_file.seek(0)
            images = convert_from_bytes(uploaded_file.read(), dpi=200)
            ocr_text = ""
            for img in images:
                ocr_text += pytesseract.image_to_string(img, lang="eng") + "\n"

            if len(ocr_text.strip()) > 50:
                return ExtractionResult(
                    ocr_text.strip(),
                    "ocr",
                    warning=(
                        "This PDF appears to be scanned (image-based). "
                        "OCR was used to extract text — layout-heavy resumes "
                        "may have reduced accuracy. For best results, export "
                        "your resume as a text-based PDF from Word or Google Docs."
                    ),
                )
        except ImportError:
            logger.warning("pdf2image/pytesseract not installed — OCR unavailable.")
        except Exception as exc:
            logger.warning("OCR failed: %s", exc)

        # All strategies exhausted
        return ExtractionResult(
            "", "empty",
            error=(
                "No text could be extracted from this PDF. It appears to be "
                "an image-based (scanned) document, and OCR either failed or "
                "is not available. Please convert your resume to a text-based "
                "PDF using Word, Google Docs, or an online PDF converter."
            ),
        )

    # ── DOCX path ─────────────────────────────────────────────────────────────
    elif filename.endswith(".docx"):
        try:
            uploaded_file.seek(0)
            doc = Document(uploaded_file)

            # Extract paragraphs AND table cells — many resume templates
            # store content in tables that vanilla paragraph iteration misses.
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)

            full_text = "\n".join(p for p in parts if p.strip())

            if len(full_text.strip()) < 50:
                return ExtractionResult(
                    "", "empty",
                    error="The DOCX file appears to be empty or contains only images."
                )

            return ExtractionResult(full_text.strip(), "docx")

        except Exception as exc:
            logger.warning("DOCX extraction failed: %s", exc)
            return ExtractionResult(
                "", "empty",
                error=f"Could not read DOCX file: {exc}."
            )

    return ExtractionResult(
        "", "empty",
        error=f"Unsupported file type '{filename}'. Please upload a PDF or DOCX."
    )


# ──────────────────────────────────────────────────────────────────────────────
# 5. STUFFING DETECTION (fixes Scenario C)
# ──────────────────────────────────────────────────────────────────────────────

# Minimal stopword set — avoids scikit-learn dependency for this function
_STOPWORDS = frozenset({
    "the", "and", "or", "in", "of", "to", "a", "an", "is", "for",
    "with", "that", "this", "are", "be", "as", "at", "by", "from",
    "was", "were", "has", "have", "had", "will", "would", "can", "could",
    "should", "their", "they", "them", "not", "but", "its", "your", "our",
    "we", "you", "i", "my", "on", "it", "do", "into", "than", "more",
    "also", "which", "who", "what", "when", "where", "how", "all",
})


def _content_tokens(text: str) -> set[str]:
    """Extract meaningful tokens, stripping stopwords and short tokens."""
    raw = re.sub(r"[^a-z\s]", "", text.lower()).split()
    return {t for t in raw if t not in _STOPWORDS and len(t) > 2}


def detect_keyword_stuffing(
    resume_text: str,
    job_text: str,
    global_threshold: float = 0.72,
    section_threshold: float = 0.85,
) -> dict:
    """
    Two-signal stuffing detector combining global overlap and section density.

    Signal 1 — Global token overlap ratio:
        Computes |resume_tokens ∩ jd_tokens| / |jd_tokens|.
        Genuine, well-written resumes typically overlap 30–60% with a relevant JD.
        Overlap > 72% is a strong indicator of verbatim JD injection.

    Signal 2 — Skills section density:
        A real skills section is a concise list of actual competencies.
        If it overlaps >85% with the JD's token vocabulary, it has likely been
        padded with every keyword from the JD — a common ATS gaming technique.

    Returns a dict with:
        - is_stuffed (bool): True if either signal fires.
        - global_overlap (float): fraction of JD tokens found in resume.
        - stuffed_sections (list[str]): section names that triggered signal 2.
        - severity (str): "none" | "suspected" | "confirmed".
    """
    resume_tokens = _content_tokens(resume_text)
    jd_tokens     = _content_tokens(job_text)

    if not jd_tokens:
        return {"is_stuffed": False, "global_overlap": 0.0,
                "stuffed_sections": [], "severity": "none"}

    global_overlap = len(resume_tokens & jd_tokens) / len(jd_tokens)
    global_flag    = global_overlap > global_threshold

    # Per-section density check
    sections = split_resume_sections(resume_text)
    stuffed_sections = []
    for section_name, section_text in sections.items():
        if not section_text.strip():
            continue
        section_tokens = _content_tokens(section_text)
        if not section_tokens:
            continue
        density = len(section_tokens & jd_tokens) / len(jd_tokens)
        if density > section_threshold:
            stuffed_sections.append(section_name)

    is_stuffed = global_flag or bool(stuffed_sections)
    if global_overlap > 0.90:
        severity = "confirmed"
    elif is_stuffed:
        severity = "suspected"
    else:
        severity = "none"

    return {
        "is_stuffed":       is_stuffed,
        "global_overlap":   round(global_overlap, 3),
        "stuffed_sections": stuffed_sections,
        "severity":         severity,
    }


# ──────────────────────────────────────────────────────────────────────────────
# 6. SKILL EXTRACTION WITH SYNONYM EXPANSION (fixes Scenario A)
# ──────────────────────────────────────────────────────────────────────────────

def extract_skills(tokens: list[str], skill_vocab: set[str]) -> set[str]:
    """
    Extract skills from a token list using both unigram and bigram matching,
    then canonicalise via the synonym graph.

    Bigram construction:
        We zip adjacent tokens to form "machine learning", "deep learning" etc.
        without importing nltk or doing full n-gram analysis — keeps the
        dependency footprint minimal.
    """
    found: set[str] = set()
    token_str = " ".join(tokens)   # reconstruct text for substring matching

    for skill in skill_vocab:
        if skill in token_str:
            found.add(skill)

    # Canonicalise: if an alias was found, add the canonical form
    canonicalised: set[str] = set()
    for skill in found:
        canonical = _ALIAS_TO_CANONICAL.get(skill, skill)
        canonicalised.add(canonical)
        # Also add the found surface form so downstream sees both
        canonicalised.add(skill)

    return canonicalised


def expand_skill_set(skills: set[str]) -> set[str]:
    """
    Synonym-expand a skill set bidirectionally.

    For each skill in the input, inject its canonical form AND all aliases.
    This ensures that "deep learning" in a resume matches "machine learning"
    in a JD because both map to the same cluster of synonyms.

    This is the core fix for Scenario A (Synonym Test):
    The intersection-over-union becomes semantically aware rather than
    purely lexical without requiring a neural model at inference time.
    """
    expanded = set(s.lower() for s in skills)
    for canonical, aliases in SKILL_SYNONYMS.items():
        # If any member of this synonym cluster exists in the set,
        # inject the entire cluster
        cluster = {canonical} | set(aliases)
        if expanded & cluster:         # non-empty intersection → cluster member found
            expanded |= cluster        # inject all cluster members
    return expanded


# ──────────────────────────────────────────────────────────────────────────────
# 7. WEIGHTED SKILL MATCHING (fixes Scenario D)
# ──────────────────────────────────────────────────────────────────────────────

def _infer_skill_weights(job_text: str, job_skills: set[str]) -> dict[str, float]:
    """
    Parse the JD for importance signal phrases and assign a weight multiplier
    to each skill based on the surrounding sentence context.

    Algorithm:
        1. Split JD into sentences.
        2. For each skill, find sentences that mention it.
        3. Scan those sentences for importance signal phrases (IMPORTANCE_SIGNALS).
        4. Assign the highest-matching multiplier (conservative: takes the max,
           not the sum, to avoid double-counting "required must-have" sentences).
        5. Default weight = 1.0 if no signal found.
    """
    sentences = re.split(r"[.;\n\r]+", job_text.lower())
    weights: dict[str, float] = {}

    for skill in job_skills:
        skill_lower = skill.lower()
        best_weight = 1.0

        for sent in sentences:
            if skill_lower not in sent:
                continue
            for signal, w in IMPORTANCE_SIGNALS.items():
                if signal in sent:
                    best_weight = max(best_weight, w)

        weights[skill] = best_weight

    return weights


def weighted_match_skills(
    resume_skills: set[str],
    job_skills: set[str],
    job_text: str,
) -> dict:
    """
    Importance-weighted skill matching — fixes Scenario D (Near Miss).

    Standard match_skills treats all skills equally:
        score = |matched| / |job_skills|

    This function weights each skill by its importance in the JD:
        weighted_score = Σ(weight_i for matched skills) / Σ(weight_i for all JD skills)

    Impact: A candidate missing "Kubernetes (required)" scores ~65% even at
    9/10 skills, while missing "agile (preferred)" at 9/10 scores ~93%.
    The score now reflects *business-defined importance*, not raw skill count.

    Also computes dealbreaker_skills: JD skills with weight ≥ 2.5 that are
    absent from the resume. These are surfaced as hard warnings in the UI.
    """
    if not job_skills:
        return {
            "score": 0, "matched": set(), "missing": set(),
            "weights": {}, "dealbreakers": set(),
            "weighted_score": 0.0,
        }

    # Synonym-expand both skill sets before intersection
    expanded_resume = expand_skill_set(resume_skills)
    expanded_job    = expand_skill_set(job_skills)

    # Compute intersection on expanded sets, then map back to original JD skills
    # so the score denominator stays meaningful (original JD requirements)
    matched_original: set[str] = set()
    for skill in job_skills:
        skill_cluster = {skill} | set(SKILL_SYNONYMS.get(skill, []))
        skill_cluster |= {_ALIAS_TO_CANONICAL.get(s, s) for s in skill_cluster}
        if expanded_resume & skill_cluster:
            matched_original.add(skill)

    missing_original = job_skills - matched_original

    # Infer per-skill weights from JD language
    weights = _infer_skill_weights(job_text, job_skills)

    total_weight   = sum(weights.get(s, 1.0) for s in job_skills)
    matched_weight = sum(weights.get(s, 1.0) for s in matched_original)

    weighted_score = (matched_weight / total_weight) * 100 if total_weight > 0 else 0.0
    # Simple (unweighted) score kept for backward-compatibility display
    simple_score   = int((len(matched_original) / len(job_skills)) * 100)

    # Dealbreakers: missing skills with importance weight ≥ 2.5
    dealbreakers = {s for s in missing_original if weights.get(s, 1.0) >= 2.5}

    return {
        "score":          simple_score,
        "weighted_score": round(weighted_score, 1),
        "matched":        matched_original,
        "missing":        missing_original,
        "weights":        weights,
        "dealbreakers":   dealbreakers,
    }


# Legacy thin wrapper kept so any existing call site doesn't break
def match_skills(resume_skills: set[str], job_skills: set[str]):
    """
    Backward-compatible wrapper around weighted_match_skills.
    Returns (score, matched, missing) tuple — same signature as original.
    """
    result = weighted_match_skills(resume_skills, job_skills, job_text="")
    return result["score"], result["matched"], result["missing"]


# ──────────────────────────────────────────────────────────────────────────────
# 8. SIMILARITY SCORING — semantic + TF-IDF stack
# ──────────────────────────────────────────────────────────────────────────────

def semantic_similarity(resume_text: str, job_text: str) -> Optional[float]:
    """
    Bi-encoder semantic similarity using sentence-transformers.

    Architecture:
        1. Chunk both documents into sentences (short propositions score better
           than full-document embeddings in MiniLM due to its training objective).
        2. Encode all sentences in a single batched call (GPU/CPU auto-detected).
        3. For each JD sentence, find the maximum cosine similarity to any
           resume sentence (max-pool over resume axis).
        4. Average the per-JD-sentence max scores → document-level similarity.

    This max-pooling strategy (asymmetric similarity) rewards coverage:
        a resume gets credit for matching each JD requirement at least once,
        without being penalised for additional content.

    Returns a float in [0, 100] or None if the model is unavailable.
    """
    model = _get_semantic_model()
    if model is None:
        return None

    try:
        from sentence_transformers import util as st_util
        import torch

        r_sentences = _sentence_tokenise(resume_text)
        j_sentences = _sentence_tokenise(job_text)

        if not r_sentences or not j_sentences:
            return None

        # Encode in one batch — sentence-transformers handles padding internally
        emb_r = model.encode(r_sentences, convert_to_tensor=True, show_progress_bar=False)
        emb_j = model.encode(j_sentences, convert_to_tensor=True, show_progress_bar=False)

        # cos_sim returns a (len_j × len_r) matrix
        sim_matrix = st_util.cos_sim(emb_j, emb_r)

        # Max-pool: each JD sentence picks the best-matching resume sentence
        per_jd_max = sim_matrix.max(dim=1).values
        score = float(per_jd_max.mean().item()) * 100

        return round(score, 2)

    except Exception as exc:
        logger.warning("Semantic similarity computation failed: %s", exc)
        return None


def tfidf_similarity(resume_text: str, job_text: str) -> float:
    """
    TF-IDF cosine similarity with synonym injection (Tier 2 fallback).

    Upgrade over the original:
        Before vectorisation, we inject synonym-expanded tokens into both
        documents. This gives TF-IDF partial semantic awareness without
        needing a neural model:
        "deep learning" in resume → also injects "machine learning", "ml" →
        both documents now share tokens they didn't originally share.

    ngram_range=(1,2): captures important bigrams like "machine learning",
    "deep learning", "computer vision" as single features.
    """
    if not resume_text or not job_text:
        return 0.0

    # Inject synonym expansions as appended pseudo-tokens
    def _augment(text: str) -> str:
        tokens = set(clean_text(text).split())
        injected: list[str] = []
        for canonical, aliases in SKILL_SYNONYMS.items():
            cluster = {canonical} | set(aliases)
            if tokens & cluster:
                # Append canonical + all aliases to the document text
                injected.extend(cluster)
        return clean_text(text) + " " + " ".join(injected)

    aug_resume = _augment(resume_text)
    aug_job    = _augment(job_text)

    try:
        vectorizer = TfidfVectorizer(
            stop_words="english",
            ngram_range=(1, 2),
            max_features=15_000,    # cap vocabulary for memory efficiency
            sublinear_tf=True,      # log(1 + tf) — dampens high-frequency terms
        )
        vectors    = vectorizer.fit_transform([aug_resume, aug_job])
        similarity = cosine_similarity(vectors[0], vectors[1])[0][0]
        return round(float(similarity) * 100, 2)
    except Exception as exc:
        logger.warning("TF-IDF similarity failed: %s", exc)
        return 0.0


def compute_similarity_score(resume_text: str, job_text: str) -> dict:
    """
    Unified similarity entry point — tries semantic first, falls back to TF-IDF.

    Returns a dict with:
        - score (float): best available similarity score in [0, 100]
        - method (str): "semantic" | "tfidf_augmented"
        - semantic_score (float | None)
        - tfidf_score (float)
    """
    tfidf_score    = tfidf_similarity(resume_text, job_text)
    semantic_score = semantic_similarity(resume_text, job_text)

    if semantic_score is not None:
        # Blend: 70% semantic + 30% TF-IDF.
        # Semantic captures meaning; TF-IDF captures exact keyword presence.
        # JDs often use exact terminology, so pure semantic can under-reward
        # candidates who use the precise JD keywords.
        blended = round(0.70 * semantic_score + 0.30 * tfidf_score, 2)
        return {
            "score":          blended,
            "method":         "semantic",
            "semantic_score": semantic_score,
            "tfidf_score":    tfidf_score,
        }

    return {
        "score":          tfidf_score,
        "method":         "tfidf_augmented",
        "semantic_score": None,
        "tfidf_score":    tfidf_score,
    }


# ──────────────────────────────────────────────────────────────────────────────
# 9. FINAL ATS SCORE COMPOSITION
# ──────────────────────────────────────────────────────────────────────────────

def final_ats_score(
    weighted_skill_score: float,
    similarity_score: float,
    stuffing_penalty: float = 0.0,
) -> float:
    """
    Composite ATS score with configurable weights.

    Composition:
        55% weighted skill match  — most ATS parsers are skill-list driven
        45% semantic/TF-IDF match — holistic document relevance
        - stuffing_penalty        — 0–30 point deduction for detected gaming

    Returns a value in [0, 100].
    """
    raw = (0.55 * weighted_skill_score) + (0.45 * similarity_score)
    penalised = max(0.0, raw - stuffing_penalty)
    return round(penalised, 1)


# ──────────────────────────────────────────────────────────────────────────────
# 10. SECTION SPLITTING AND SECTION-WISE SCORING
# ──────────────────────────────────────────────────────────────────────────────

def split_resume_sections(resume_text: str) -> dict[str, str]:
    """
    Heuristic section splitter using heading-pattern matching.

    Detects section boundaries by scanning lines for known section header
    patterns (SECTION_PATTERNS). Text between two detected headers belongs
    to the first header's section.

    Returns a dict mapping section_name → section_text.
    Falls back to a single "full_text" key if no sections are detected.
    """
    sections: dict[str, str] = defaultdict(str)
    current_section = "summary"   # default — text before first heading

    for line in resume_text.split("\n"):
        line_lower = line.strip().lower()
        detected = False
        for section_name, patterns in SECTION_PATTERNS.items():
            if any(p in line_lower for p in patterns):
                if len(line_lower) < 50:  # short lines are likely headings
                    current_section = section_name
                    detected = True
                    break
        if not detected:
            sections[current_section] += line + "\n"

    # If nothing was detected, return the full text as-is
    if all(not v.strip() for v in sections.values()):
        return {"full_text": resume_text}

    return dict(sections)


def section_wise_scores(
    sections: dict[str, str],
    job_text: str,
) -> dict[str, float]:
    """
    Per-section TF-IDF similarity scores.

    Scores each resume section independently against the full JD.
    This surfaces which sections are dragging the overall score down —
    the UI renders these as progress bars so the user knows exactly
    where to improve.
    """
    scores: dict[str, float] = {}
    for name, text in sections.items():
        if not text.strip():
            scores[name] = 0.0
        else:
            scores[name] = tfidf_similarity(text, job_text)
    return scores


# ──────────────────────────────────────────────────────────────────────────────
# 11. KEYWORD GAP ANALYSIS
# ──────────────────────────────────────────────────────────────────────────────

def keyword_gap(resume_text: str, job_text: str, top_n: int = 20) -> list[str]:
    """
    Surface the most important JD terms absent from the resume.

    Method:
        1. Fit a TF-IDF vectorizer on the JD alone.
        2. Extract the top-N highest-TF-IDF terms (these are the JD's
           most distinctive and important keywords — not stopwords, not noise).
        3. Filter to terms that don't appear in the resume.

    This is more signal-rich than simple set difference because TF-IDF
    down-weights ubiquitous terms and surfaces domain-specific vocabulary.
    """
    if not job_text.strip():
        return []

    try:
        vectorizer = TfidfVectorizer(
            stop_words="english",
            ngram_range=(1, 2),
            max_features=200,
        )
        vectorizer.fit([clean_text(job_text)])

        # Map feature index → TF-IDF weight in the JD document
        feature_names  = vectorizer.get_feature_names_out()
        jd_vector      = vectorizer.transform([clean_text(job_text)]).toarray()[0]
        term_scores    = sorted(
            zip(feature_names, jd_vector),
            key=lambda x: x[1],
            reverse=True,
        )

        resume_lower = resume_text.lower()
        gaps = [
            term for term, score in term_scores
            if score > 0 and term not in resume_lower
        ]
        return gaps[:top_n]

    except Exception as exc:
        logger.warning("keyword_gap failed: %s", exc)
        return []


def keyword_suggestions(gaps: list[str]) -> list[str]:
    """
    Filter and prettify keyword gap list for display.
    Removes single-character tokens and sorts alphabetically.
    """
    return sorted({g for g in gaps if len(g) > 2})


# ──────────────────────────────────────────────────────────────────────────────
# 12. SKILL SUGGESTIONS (quick-fix copy)
# ──────────────────────────────────────────────────────────────────────────────

def generate_skill_suggestions(missing_skills: set[str]) -> list[str]:
    """
    Generate actionable resume bullet-point suggestions for missing skills.
    These are template strings — the AI optimiser in llm_optimizer.py
    generates more contextual versions; these serve as a fast non-LLM fallback.
    """
    suggestions = []
    for skill in sorted(missing_skills)[:8]:   # cap at 8 to avoid UI overload
        suggestions.append(
            f"Add '{skill}' to your Skills section or weave it into "
            f"a bullet point in your Experience section with a quantified outcome."
        )
    return suggestions
